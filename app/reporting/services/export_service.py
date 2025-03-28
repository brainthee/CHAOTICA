import io
import csv
import datetime
import json
from django.http import HttpResponse, FileResponse
from django.template.loader import render_to_string
from django.utils.text import slugify

class ExportService:
    """
    Service for exporting report data to various formats
    """
    
    @staticmethod
    def export_report(report, data, format_type=None):
        """
        Export report data to the specified format
        """
        # If format_type is not specified, use the report's preferred format
        if format_type is None:
            format_type = report.presentation_type
            
        # Call the appropriate export method based on format_type
        export_methods = {
            'excel': ExportService.export_to_excel,
            'pdf': ExportService.export_to_pdf,
            'html': ExportService.export_to_html,
            'csv': ExportService.export_to_csv,
            'word': ExportService.export_to_word,
            'text': ExportService.export_to_text,
        }
        
        # Use the selected export method or default to CSV if not found
        export_method = export_methods.get(format_type, ExportService.export_to_csv)
        
        # Generate filename
        filename = f"{slugify(report.name)}_{datetime.datetime.now().strftime('%Y%m%d_%H%M%S')}"
        
        # Get the fields information
        fields = report.get_fields()
        field_names = [field.get_display_name() for field in fields]
        
        # Call the export method
        return export_method(data, field_names, filename, report)
    
    @staticmethod
    def export_to_excel(data, field_names, filename, report=None):
        """
        Export data to Excel format
        """
        try:
            import xlsxwriter
        except ImportError:
            # Fallback to CSV if xlsxwriter is not available
            return ExportService.export_to_csv(data, field_names, filename, report)
        
        # Create an in-memory output file for the Excel workbook
        output = io.BytesIO()
        
        # Create Excel workbook and worksheet
        workbook = xlsxwriter.Workbook(output, {'remove_timezone': True})
        worksheet = workbook.add_worksheet('Report')
        
        # Add a header format
        header_format = workbook.add_format({
            'bold': True,
            'border': 1,
            'bg_color': '#D3D3D3',  # Light gray background
            'text_wrap': True,
            'valign': 'vcenter',
            'align': 'center',
        })
        
        # Write the headers
        for col_num, column_title in enumerate(field_names):
            worksheet.write(0, col_num, column_title, header_format)
            worksheet.set_column(col_num, col_num, 15)  # Set column width
        
        # Write the data rows
        row_num = 1
        for row in data:
            for col_num, cell_value in enumerate(row.values()):
                worksheet.write(row_num, col_num, cell_value)
            row_num += 1
        
        # Set auto-filter
        worksheet.autofilter(0, 0, row_num - 1, len(field_names) - 1)
        
        # Close the workbook
        workbook.close()
        
        # Seek to the beginning and create response
        output.seek(0)
        
        # Create the HttpResponse object with Excel headers
        response = HttpResponse(
            output.getvalue(),
            content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet'
        )
        response['Content-Disposition'] = f'attachment; filename="{filename}.xlsx"'
        
        return response
    
    @staticmethod
    def export_to_pdf(data, field_names, filename, report=None):
        """
        Export data to PDF format
        """
        try:
            from reportlab.lib.pagesizes import letter, landscape
            from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph
            from reportlab.lib.styles import getSampleStyleSheet
            from reportlab.lib import colors
        except ImportError:
            # Fallback to CSV if reportlab is not available
            return ExportService.export_to_csv(data, field_names, filename, report)
        
        # Create an in-memory output file for the PDF
        buffer = io.BytesIO()
        
        # Create the PDF document with landscape orientation
        doc = SimpleDocTemplate(buffer, pagesize=landscape(letter))
        
        # Get styles
        styles = getSampleStyleSheet()
        title_style = styles['Heading1']
        
        # Create the content for the PDF
        elements = []
        
        # Add report title
        if report:
            elements.append(Paragraph(report.name, title_style))
        
        # Prepare the data for the table
        table_data = [field_names]  # Headers
        
        # Add the data rows
        for row in data:
            table_data.append([str(value) for value in row.values()])
        
        # Create the table
        table = Table(table_data)
        
        # Style the table
        table_style = TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
            ('ALIGN', (0, 0), (-1, 0), 'CENTER'),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
            ('BACKGROUND', (0, 1), (-1, -1), colors.white),
            ('GRID', (0, 0), (-1, -1), 1, colors.black),
        ])
        
        # Apply the style
        table.setStyle(table_style)
        
        # Add the table to the elements
        elements.append(table)
        
        # Build the PDF
        doc.build(elements)
        
        # Get the PDF data from the buffer
        pdf_data = buffer.getvalue()
        buffer.close()
        
        # Create the response
        response = HttpResponse(content_type='application/pdf')
        response['Content-Disposition'] = f'attachment; filename="{filename}.pdf"'
        response.write(pdf_data)
        
        return response
    
    @staticmethod
    def export_to_html(data, field_names, filename, report=None):
        """
        Export data to HTML format
        """
        # Prepare context for the template
        context = {
            'report': report,
            'field_names': field_names,
            'data': data,
            'title': report.name if report else 'Report',
            'generated_at': datetime.datetime.now(),
        }
        
        # Render the HTML content
        html_content = render_to_string('reporting/exports/report_html.html', context)
        
        # Create the response
        response = HttpResponse(content_type='text/html')
        response['Content-Disposition'] = f'attachment; filename="{filename}.html"'
        response.write(html_content)
        
        return response
    
    @staticmethod
    def export_to_csv(data, field_names, filename, report=None):
        """
        Export data to CSV format
        """
        # Create a file-like buffer to receive CSV data
        buffer = io.StringIO()
        
        # Create CSV writer
        writer = csv.writer(buffer)
        
        # Write headers
        writer.writerow(field_names)
        
        # Write data rows
        for row in data:
            writer.writerow(row.values())
        
        # Create the HTTP response
        response = HttpResponse(buffer.getvalue(), content_type='text/csv')
        response['Content-Disposition'] = f'attachment; filename="{filename}.csv"'
        
        return response
    
    @staticmethod
    def export_to_word(data, field_names, filename, report=None):
        """
        Export data to Word format (simple HTML that Word can open)
        """
        try:
            from docx import Document
            from docx.shared import Inches
        except ImportError:
            # Fallback to HTML if python-docx is not available
            return ExportService.export_to_html(data, field_names, filename, report)
        
        # Create a new document
        document = Document()
        
        # Add report title
        if report:
            document.add_heading(report.name, level=1)
        
        # Add generated timestamp
        document.add_paragraph(f"Generated: {datetime.datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
        
        # Add a table
        table = document.add_table(rows=1, cols=len(field_names))
        table.style = 'Table Grid'
        
        # Add header row
        header_cells = table.rows[0].cells
        for i, name in enumerate(field_names):
            header_cells[i].text = name
        
        # Add data rows
        for row_data in data:
            row_cells = table.add_row().cells
            for i, value in enumerate(row_data.values()):
                row_cells[i].text = str(value)
        
        # Save to a BytesIO object
        buffer = io.BytesIO()
        document.save(buffer)
        buffer.seek(0)
        
        # Create response
        response = HttpResponse(
            buffer.getvalue(),
            content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
        response['Content-Disposition'] = f'attachment; filename="{filename}.docx"'
        
        return response
    
    @staticmethod
    def export_to_text(data, field_names, filename, report=None):
        """
        Export data to plain text format
        """
        # Create a file-like buffer
        buffer = io.StringIO()
        
        # Write report title if available
        if report:
            buffer.write(f"{report.name}\n")
            buffer.write('=' * len(report.name) + '\n\n')
        
        # Write generated timestamp
        buffer.write(f"Generated: {datetime.datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n\n")
        
        # Determine column widths
        col_widths = [len(name) for name in field_names]
        for row in data:
            for i, value in enumerate(row.values()):
                col_widths[i] = max(col_widths[i], len(str(value)))
        
        # Write headers
        header_row = ' | '.join(name.ljust(col_widths[i]) for i, name in enumerate(field_names))
        buffer.write(header_row + '\n')
        buffer.write('-' * len(header_row) + '\n')
        
        # Write data rows
        for row in data:
            row_str = ' | '.join(str(value).ljust(col_widths[i]) for i, value in enumerate(row.values()))
            buffer.write(row_str + '\n')
        
        # Create response
        response = HttpResponse(buffer.getvalue(), content_type='text/plain')
        response['Content-Disposition'] = f'attachment; filename="{filename}.txt"'
        
        return response